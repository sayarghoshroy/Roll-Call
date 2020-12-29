# Read and Process Raw Data from a `docx` file

from docx import Document
document = Document('./sample/input.docx')

all_names = []

for paragraph in document.paragraphs:
  # Process each paragraph unit separately
  unit = paragraph.text.split(',')
  
  # Handle empty units
  if unit == [] or unit == ['']:
    continue
  
  for name in unit:
    if name.strip() != '':
      all_names.append(name.strip())

print('Number of Names Extracted: ' + str(len(all_names)))
all_names = list(set(all_names))

print('Number of Unique Names: ' + str(len(all_names)))

# Sort Unicode Strings
all_names.sort()

# View Extracted Names
limit = 5
count = 1

for unit in all_names:
  if count > limit:
    break
  print(str(count) +". " +str(unit))
  count += 1

# Create Output Doc
out = Document()

output_string = ''
size = len(all_names)

for index, elem in enumerate(all_names):
  if index == (size - 1):
    output_string = output_string + elem
  else:
    output_string = output_string + elem + ', '

out.add_paragraph(str(output_string))
out.save('./sample/output.docx')

# Thank You